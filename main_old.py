import docx
import os
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from googleapiclient.discovery import build
from google.oauth2.service_account import Credentials

# Загрузка данных из Google Sheets
def load_google_sheet(s_id, s_range):
    creds = Credentials.from_service_account_file('service.json')
    service = build('sheets', 'v4', credentials=creds)
    sheet = service.spreadsheets()
    result = sheet.values().get(spreadsheetId=s_id, range=s_range).execute()
    return result.get('values', [])

# Установка стиля для документа
def set_document_style(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.5
    paragraph_format.first_line_indent = Cm(1)
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.sections[0].left_margin = Cm(2)
    doc.sections[0].right_margin = Cm(2)
    doc.sections[0].top_margin = Cm(2)
    doc.sections[0].bottom_margin = Cm(2)

def convert_to_initials(full_name):
    parts = full_name.split()
    if len(parts) == 3:
        surname, name, patronymic = parts
        return f"{surname} {name[0]}.{patronymic[0]}."
    elif len(parts) == 2:
        surname, name = parts
        return f"{surname} {name[0]}."
    else:
        return full_name

def generate_conference_program(student_data, tech_data):
    doc = docx.Document()
    set_document_style(doc)

    first_paragraph = doc.add_paragraph(
        'Форма представления материалов для программы 78 МСНК ГУАП',
        style='Normal'
    )
    first_paragraph.runs[0].bold = True
    first_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Секция кафедры
    section_heading = doc.add_paragraph()
    run1 = section_heading.add_run(f" {' ' * 4} Секция каф. ")
    run1.bold = True
    run1.italic = True
    run2 = section_heading.add_run(f"{tech_data[0][0]}. {tech_data[0][1]}")
    run2.bold = True
    run2.italic = True

    # Научный руководитель
    doc.add_paragraph(
        f" {' ' * 10} Научный руководитель секции - {tech_data[0][2]}",
        style='Normal'
    )

    doc.add_paragraph(
        f" {' ' * 10} {tech_data[0][3]}",
        style='Normal'
    )

    # Зам. научного руководителя
    doc.add_paragraph(
        f" {' ' * 10} Зам. научного руководителя секции - {tech_data[0][6]}",
        style='Normal'
    )

    doc.add_paragraph(
        f" {' ' * 10} {tech_data[0][7]}",
        style='Normal'
    )
    
    max_value = max([int(row[8]) for row in student_data if row[8].isdigit()])
    
    for cur_num in range(1, max_value + 1):
        
        # Заседание
        session_heading = doc.add_paragraph(f'Заседание {str(cur_num)}', style='Normal')
        session_heading.runs[0].bold = True
        
        # Дата и время + адрес
        session_info = doc.add_paragraph()
        session_info.add_run(f"{tech_data[cur_num - 1][11]}{' ' * 35}Санкт-Петербург, ул. Большая Морская, д. 67,")

        room_info = doc.add_paragraph()
        run = room_info.add_run(f"{' ' * 75} лит. А, ауд. {tech_data[cur_num - 1][12]}")

        # Список участников с темами
        participant_num = 1
        for row in student_data:
            if int(row[8]) == cur_num:
                initials = convert_to_initials(row[1]) 
                doc.add_paragraph(f'{participant_num}. {initials}', style='Normal')
                doc.add_paragraph(f'{row[2]}', style='Normal')
                participant_num += 1
                
    doc.save('report/Программа конференции.docx')

def generate_conference_report(student_data, tech_data):
    doc = docx.Document()
    set_document_style(doc)

    first_paragraph = doc.add_paragraph(
        'Отчёт о конференции 78 МСНК ГУАП',
        style='Normal'
    )
    first_paragraph.runs[0].bold = True
    first_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Секция кафедры
    section_heading = doc.add_paragraph()
    run1 = section_heading.add_run(f" {' ' * 4} Секция каф. ")
    run1.bold = True
    run1.italic = True
    run2 = section_heading.add_run(f"{tech_data[0][0]}. {tech_data[0][1]}")
    run2.bold = True
    run2.italic = True

    max_value = max([int(row[8]) for row in student_data if row[8].isdigit()])
    
    for cur_num in range(1, max_value + 1):
        
        session_heading = doc.add_paragraph(f'Заседание {str(cur_num)}', style='Normal')
        session_heading.runs[0].bold = True
        
        session_info = doc.add_paragraph()
        session_info.add_run(f"{tech_data[cur_num - 1][11]}{' ' * 35}Санкт-Петербург, ул. Большая Морская, д. 67,")

        room_info = doc.add_paragraph()
        run = room_info.add_run(f"{' ' * 75} лит. А, ауд. {tech_data[cur_num - 1][12]}")

        doc.add_paragraph(
            f"Научный руководитель секции - {tech_data[0][3]} {convert_to_initials(tech_data[0][2])}",
            style='Normal'
        )

        doc.add_paragraph(
            f"Список докладов",
            style='Normal'
        )

        # Таблица для списка докладов
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '№ п/п'
        hdr_cells[1].text = 'ФИО докладчика, название доклада'
        hdr_cells[2].text = 'Статус (магистр/студент)'
        hdr_cells[3].text = 'Решение'

        # Выравнивание текста в заголовке таблицы по центру
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.first_line_indent = Cm(0)

        # Заполнение таблицы
        participant_num = 1
        for row in student_data:
            if int(row[8]) == cur_num:
                initials = row[1]
                status = f"{row[4]} Гр. № {row[5]}"
                recommendation = "опубликовать доклад в сборнике МСНК" if row[9] == "1" else "доклад плохо подготовлен"

                row_cells = table.add_row().cells
                row_cells[0].text = str(participant_num)
                row_cells[1].text = f"{initials}\n{row[2]}" 
                row_cells[2].text = status 
                row_cells[3].text = recommendation  

                for paragraph in row_cells[1].paragraphs + row_cells[2].paragraphs + row_cells[3].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    paragraph.paragraph_format.first_line_indent = Cm(0)

                participant_num += 1

        doc.add_paragraph()

    doc.add_paragraph("Подпись научного руководителя секции", style='Normal')

    doc.save('report/Отчёт о конференции.docx')


def generate_conference_list(student_data, tech_data):
    doc = docx.Document()
    set_document_style(doc)

    first_paragraph = doc.add_paragraph(
        'Список представляемых к публикации докладов',
        style='Normal'
    )
    first_paragraph.runs[0].bold = True
    first_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Секция кафедры
    doc.add_paragraph(f"Кафедра {tech_data[0][1]}")
    doc.add_paragraph(tech_data[0][2])
    doc.add_paragraph(f"e-mail: {tech_data[0][4]}")
    doc.add_paragraph(f"тел.: {tech_data[0][5]}")

    # Пополнение списка студентов
    for row in student_data:
        if row[9] == "1":
            combined_paragraph = doc.add_paragraph()
            
            name_run = combined_paragraph.add_run(f"{convert_to_initials(row[1])} ")
            name_run.italic = True
            
            combined_paragraph.add_run(f"{row[2]}")

    doc.add_paragraph("\n" * 2)
    doc.add_paragraph(f"Руководитель УНИДС {' ' * 40}{convert_to_initials(tech_data[0][2])}")


    doc.save('report/Список представляемых к публикации докладов.docx')

if __name__ == "__main__":
    if not os.path.exists('report'):
        os.makedirs('report')

    student_sheet_id = '1szrnHnYN1LOLG9V8eJeye3YTZY4Ka_FPPZfWjQ5_zlw'
    tech_sheet_id = '1szrnHnYN1LOLG9V8eJeye3YTZY4Ka_FPPZfWjQ5_zlw'
    student_range = 'Sheet1!A2:L' 
    tech_range = 'Sheet2!A2:M'
    
    student_data = load_google_sheet(student_sheet_id, student_range)
    tech_data = load_google_sheet(tech_sheet_id, tech_range)
    
    # CLI для выбора типа документа
    print("Какой документ хотите составить?")
    print("1. Программа конференции")
    print("2. Отчёт о конференции")
    print("3. Список представляемых к публикации докладов")
    print("0. Выйти")
    while True:
        document_type = input("Введите номер документа (1, 2 или 3): ")
        if document_type == '1':
            generate_conference_program(student_data, tech_data)
            print("Сгенерирована программа конференции.")
        elif document_type == '2':
            generate_conference_report(student_data, tech_data)
            print("Сгенерирован отчет о конференции.")
        elif document_type == '3':
            generate_conference_list(student_data, tech_data)
            print("Сгенерирован список представляемых к публикации докладов")
        elif document_type == '0':
            print("Завершение программы")
            break
        else:
            print("Ошибка: Неправильный формат ввода.")


