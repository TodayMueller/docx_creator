from fastapi import FastAPI, HTTPException, Depends
from fastapi.responses import FileResponse
from googleapiclient.discovery import build
from google.oauth2.service_account import Credentials
from typing import Optional, List, Tuple
from datetime import datetime
import os
import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import logging

app = FastAPI()

GOOGLE_SHEET_ID = "1wuXlICCDQyWFyiOnjh1TDxfO0EJPn6deAYGwXLF_O54"
TECH_RANGE = "Sheet1!A2:S"

# Load data from Google Sheets
def load_google_sheet(s_id: str, s_range: str) -> List[List[str]]:
    try:
        creds = Credentials.from_service_account_file("service.json")
        service = build("sheets", "v4", credentials=creds)
        sheet = service.spreadsheets()
        result = sheet.values().get(spreadsheetId=s_id, range=s_range).execute()
        logging.info(f"Google Sheets data: {result}")
        return result.get("values", [])
    except Exception as e:
        logging.exception(f"Error loading data from Google Sheets: {e}")
        raise HTTPException(status_code=500, detail=f"Error loading data from Google Sheets: {e}")

# Set document style for Word documents
def set_document_style(doc: docx.Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.5
    paragraph_format.first_line_indent = Cm(1)
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.sections[0].left_margin = Cm(2)
    doc.sections[0].right_margin = Cm(2)
    doc.sections[0].top_margin = Cm(2)
    doc.sections[0].bottom_margin = Cm(2)

# Convert full name to initials
def convert_to_initials(full_name: str) -> str:
    parts = full_name.split()
    if len(parts) == 3:
        surname, name, patronymic = parts
        return f"{surname} {name[0]}.{patronymic[0]}."
    elif len(parts) == 2:
        surname, name = parts
        return f"{surname} {name[0]}."
    else:
        return full_name

# Format date as "DD Month YYYY"
def format_date(date_str: str) -> str:
    months = {
        "01": "января",
        "02": "февраля",
        "03": "марта",
        "04": "апреля",
        "05": "мая",
        "06": "июня",
        "07": "июля",
        "08": "августа",
        "09": "сентября",
        "10": "октября",
        "11": "ноября",
        "12": "декабря",
    }
    date = datetime.strptime(date_str, "%Y-%m-%d")
    return f"{date.day} {months[date.strftime('%m')]} {date.year}г."

# Generate conference program document
def generate_conference_program(tech_data: List[List[str]]) -> Path:
    doc = docx.Document()
    set_document_style(doc)

    first_paragraph = doc.add_paragraph("Форма представления материалов для программы 78 МСНК ГУАП", style="Normal")
    first_paragraph.runs[0].bold = True
    first_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    section_heading = doc.add_paragraph("Секция каф. 43. Компьютерных технологий и программной инженерии")
    run1 = section_heading.add_run()
    run1.bold = True
    run1.italic = True

    # Sort tech_data by date and create sessions
    tech_data.sort(key=lambda x: datetime.strptime(x[15], "%Y-%m-%d"))
    sessions = {}
    session_num = 1
    for row in tech_data:
        date = row[15]
        if date not in sessions:
            sessions[date] = session_num
            session_num += 1

    # Generate program content
    for date, session_id in sessions.items():
        session_heading = doc.add_paragraph(f"Заседание {str(session_id)}", style="Normal")
        session_info = doc.add_paragraph(format_date(date))

        participant_num = 1
        for row in tech_data:
            if sessions[row[15]] == session_id:
                initials = convert_to_initials(row[7] + " " + row[8] + " " + row[9])
                doc.add_paragraph(f"{participant_num}. {initials}")
                doc.add_paragraph(row[13])
                participant_num += 1

    file_path = Path("report/programme.docx")
    doc.save(file_path)
    return file_path

# Generate conference report document
def generate_conference_report(tech_data: List[List[str]]) -> Path:
    doc = docx.Document()
    set_document_style(doc)

    first_paragraph = doc.add_paragraph("Отчёт о конференции 78 МСНК ГУАП", style="Normal")
    first_paragraph.runs[0].bold = True
    first_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    section_heading = doc.add_paragraph("Секция каф. 43. Компьютерных технологий и программной инженерии ")
    run1 = section_heading.add_run()
    run1.bold = True
    run1.italic = True

    # Sort tech_data by date and create sessions
    tech_data.sort(key=lambda x: datetime.strptime(x[15], "%Y-%m-%d"))
    sessions = {}
    session_num = 1
    for row in tech_data:
        date = row[15]
        if date not in sessions:
            sessions[date] = session_num
            session_num += 1

    # Generate report content
    for date, session_id in sessions.items():
        session_heading = doc.add_paragraph(f"Заседание {str(session_id)}", style="Normal")
        session_info = doc.add_paragraph(format_date(date))
        doc.add_paragraph("Список докладов", style="Normal")

        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "№ п/п"
        hdr_cells[1].text = "ФИО докладчика, название доклада"
        hdr_cells[2].text = "Статус (магистр/студент)"
        hdr_cells[3].text = "Решение"

        participant_num = 1
        for row in tech_data:
            if sessions[row[15]] == session_id:
                initials = convert_to_initials(row[7] + " " + row[8] + " " + row[9])
                status = f"{row[12]} Гр. № {row[11]}" if row[11] else row[12]
                recommendation = "нет данных"

                if len(row) > 16:
                    if row[16] == "1":
                        recommendation = "опубликовать доклад в сборнике МСНК"
                    elif row[16] == "2":
                        recommendation = "опубликовать доклад в сборнике МСНК; рекомендовать к участию в финале конкурса на лучшую студенческую научную работу ГУАП"
                    elif row[16] == "0":
                        recommendation = "доклад плохо подготовлен"

                row_cells = table.add_row().cells
                row_cells[0].text = str(participant_num)
                row_cells[1].text = f"{initials}\n{row[13]}"
                row_cells[2].text = status
                row_cells[3].text = recommendation

                participant_num += 1

    file_path = Path("report/report.docx")
    doc.save(file_path)
    return file_path

# Generate conference publications list document
def generate_conference_list(tech_data: List[List[str]]) -> Path:
    doc = docx.Document()
    set_document_style(doc)

    first_paragraph = doc.add_paragraph("Список представляемых к публикации докладов", style="Normal")
    first_paragraph.runs[0].bold = True
    first_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("Кафедра 43. Компьютерных технологий и программной инженерии")

    for row in tech_data:
        if len(row) > 16 and row[16] in ["1", "2"]:
            combined_paragraph = doc.add_paragraph()
            name_run = combined_paragraph.add_run(convert_to_initials(row[7] + " " + row[8] + " " + row[9]))
            name_run.italic = True
            combined_paragraph.add_run(row[13])

    file_path = Path("report/publications.docx")
    doc.save(file_path)
    return file_path

# Endpoint for generating conference programme document
@app.get("/conferences/programme")
def get_programme() -> FileResponse:
    tech_data = load_google_sheet(GOOGLE_SHEET_ID, TECH_RANGE)
    if not tech_data:
        raise HTTPException(status_code=404, detail="Conference data not found")

    file_path = generate_conference_program(tech_data)
    return FileResponse(file_path, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename="conference_programme.docx")

# Endpoint for generating conference report document
@app.get("/conferences/report")
def get_report() -> FileResponse:
    tech_data = load_google_sheet(GOOGLE_SHEET_ID, TECH_RANGE)
    if not tech_data:
        raise HTTPException(status_code=404, detail="Conference data not found")

    file_path = generate_conference_report(tech_data)
    return FileResponse(file_path, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename="conference_report.docx")

# Endpoint for generating conference publications list document
@app.get("/conferences/publications")
def get_publications() -> FileResponse:
    tech_data = load_google_sheet(GOOGLE_SHEET_ID, TECH_RANGE)
    if not tech_data:
        raise HTTPException(status_code=404, detail="Conference data not found")

    file_path = generate_conference_list(tech_data)
    return FileResponse(file_path, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename="conference_publications.docx")

if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="127.0.0.1", port=8000)



#uvicorn v3:app --reload --log-level info
#