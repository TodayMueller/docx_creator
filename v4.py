from fastapi import FastAPI, HTTPException, Depends
from fastapi.responses import FileResponse
from googleapiclient.discovery import build
from google.oauth2.service_account import Credentials
from typing import Optional, List, Tuple
from datetime import datetime
import os
import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import logging

app = FastAPI()

GOOGLE_SHEET_ID = '1MROr3Pw7nMG2vYW_AeqIy2q9FTF7URD3b24tyrBYWgE'
STUD_RANGE = 'Sheet1!A2:S' 
TECH_RANGE = 'Sheet2!A2:N'

# Загрузка данных из Google Sheets
def load_google_sheet(s_id: str, s_range: str) -> List[List[str]]:
    try:
        creds = Credentials.from_service_account_file("service.json")
        service = build("sheets", "v4", credentials=creds)
        sheet = service.spreadsheets()
        result = sheet.values().get(spreadsheetId=s_id, range=s_range).execute()
        logging.info(f"Google Sheets data: {result}")
        return result.get("values", [])
    except Exception as e:
        logging.exception(f"Error loading data from Google Sheets: {e}")
        raise HTTPException(status_code=500, detail=f"Error loading data from Google Sheets: {e}")


# Установка стиля для документа
def set_document_style(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.5
    paragraph_format.first_line_indent = Cm(1)
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.sections[0].left_margin = Cm(2)
    doc.sections[0].right_margin = Cm(2)
    doc.sections[0].top_margin = Cm(2)
    doc.sections[0].bottom_margin = Cm(2)

def convert_to_initials(full_name):
    parts = full_name.split()
    if len(parts) == 3:
        surname, name, patronymic = parts
        return f"{surname} {name[0]}.{patronymic[0]}."
    elif len(parts) == 2:
        surname, name = parts
        return f"{surname} {name[0]}."
    else:
        return full_name
    

def format_date(date_str):
    months = {
        '01': 'января', '02': 'февраля', '03': 'марта', '04': 'апреля',
        '05': 'мая', '06': 'июня', '07': 'июля', '08': 'августа',
        '09': 'сентября', '10': 'октября', '11': 'ноября', '12': 'декабря'
    }
    date = datetime.strptime(date_str, "%Y-%m-%d")
    return f"{date.day} {months[date.strftime('%m')]} {date.year}г."

def generate_conference_program(student_data, tech_data):
    doc = docx.Document()
    set_document_style(doc)

    first_paragraph = doc.add_paragraph(
        'Форма представления материалов для программы 78 МСНК ГУАП',
        style='Normal'
    )
    first_paragraph.runs[0].bold = True
    first_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Секция кафедры
    section_heading = doc.add_paragraph()
    run1 = section_heading.add_run(f" {' ' * 4} Секция каф. ")
    run1.bold = True
    run1.italic = True
    run2 = section_heading.add_run(f"{tech_data[0][0]}. {tech_data[0][1]}")
    run2.bold = True
    run2.italic = True

    # Научный руководитель
    doc.add_paragraph(
        f" {' ' * 10} Научный руководитель секции - {tech_data[0][2]}",
        style='Normal'
    )

    doc.add_paragraph(
        f" {' ' * 10} {tech_data[0][3]}",
        style='Normal'
    )

    # Зам. научного руководителя
    doc.add_paragraph(
        f" {' ' * 10} Зам. научного руководителя секции - {tech_data[0][6]}",
        style='Normal'
    )

    doc.add_paragraph(
        f" {' ' * 10} {tech_data[0][7]}",
        style='Normal'
    )
    
    max_value = max([int(row[15]) for row in student_data if row[15].isdigit()])
    
    for cur_num in range(1, max_value + 1):
        
        # Заседание
        session_heading = doc.add_paragraph(f'Заседание {str(cur_num)}', style='Normal')
        session_heading.runs[0].bold = True
        
        # Дата и время + адрес
        session_info = doc.add_paragraph()
        # Форматируем дату и текст до нужной длины 61 символ
        formatted_text = f"{format_date(tech_data[cur_num - 1][11])}, {tech_data[cur_num - 1][12]}"
        formatted_text = formatted_text.ljust(58)  # Дополняем пробелами до 61 символа

        # Добавляем к строке адрес
        session_info.add_run(f"{formatted_text}Санкт-Петербург, ул. Большая Морская, д. 67,")

        room_info = doc.add_paragraph()
        run = room_info.add_run(f"{' ' * 73} лит. А, ауд. {tech_data[cur_num - 1][13]}")

        # Список участников с темами
        participant_num = 1
        for row in student_data:
            if int(row[15]) == cur_num:
                initials = convert_to_initials(row[7] + ' ' + row[8] + ' ' + row[9]) 
                doc.add_paragraph(f'{participant_num}. {initials}', style='Normal')
                doc.add_paragraph(f'{row[13]}', style='Normal')
                participant_num += 1
                
    file_path = 'report/programme.docx'
    doc.save(file_path)
    return file_path

def generate_conference_report(student_data, tech_data):
    doc = docx.Document()
    set_document_style(doc)

    first_paragraph = doc.add_paragraph(
        'Отчёт о конференции 78 МСНК ГУАП',
        style='Normal'
    )
    first_paragraph.runs[0].bold = True
    first_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Секция кафедры
    section_heading = doc.add_paragraph()
    run1 = section_heading.add_run(f" {' ' * 4} Секция каф. ")
    run1.bold = True
    run1.italic = True
    run2 = section_heading.add_run(f"{tech_data[0][0]}. {tech_data[0][1]}")
    run2.bold = True
    run2.italic = True

    max_value = max([int(row[15]) for row in student_data if row[15].isdigit()])
    
    for cur_num in range(1, max_value + 1):
        
        session_heading = doc.add_paragraph(f'Заседание {str(cur_num)}', style='Normal')
        session_heading.runs[0].bold = True
        
        session_info = doc.add_paragraph()
        session_info.add_run(f"{format_date(tech_data[cur_num - 1][11])}, {tech_data[cur_num - 1][12]}{' ' * 35}Санкт-Петербург, ул. Большая Морская, д. 67,")

        room_info = doc.add_paragraph()
        run = room_info.add_run(f"{' ' * 73} лит. А, ауд. {tech_data[cur_num - 1][13]}")

        doc.add_paragraph(
            f"Научный руководитель секции - {tech_data[0][3]} {convert_to_initials(tech_data[0][2])}",
            style='Normal'
        )

        doc.add_paragraph(
            f"Список докладов",
            style='Normal'
        )

        # Таблица для списка докладов
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '№ п/п'
        hdr_cells[1].text = 'ФИО докладчика, название доклада'
        hdr_cells[2].text = 'Статус (магистр/студент)'
        hdr_cells[3].text = 'Решение'

        # Выравнивание текста в заголовке таблицы по центру
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.first_line_indent = Cm(0)

        # Заполнение таблицы
        participant_num = 1
        for row in student_data:
            if int(row[15]) == cur_num:
                initials = row[7] + " " + row[8] + " " + row[9]
                status = f"{row[12]} Гр. № {row[11]}" if row[11] else row[12]
                if len(row) > 16:
                    if row[16] == "1":
                        recommendation = "опубликовать доклад в сборнике МСНК"
                    elif row[16] == "2":
                        recommendation = ("опубликовать доклад в сборнике МСНК; "
                                          "рекомендовать к участию в финале конкурса "
                                          "на лучшую студенческую научную работу ГУАП")
                    elif row[16] == "0":
                        recommendation = "доклад плохо подготовлен"
                else:
                    recommendation = "нет данных"

                row_cells = table.add_row().cells
                row_cells[0].text = str(participant_num)
                row_cells[1].text = f"{initials}\n{row[13]}" 
                row_cells[2].text = status 
                row_cells[3].text = recommendation  

                for paragraph in row_cells[1].paragraphs + row_cells[2].paragraphs + row_cells[3].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    paragraph.paragraph_format.first_line_indent = Cm(0)

                participant_num += 1

        doc.add_paragraph()

    doc.add_paragraph("Подпись научного руководителя секции", style='Normal')

    file_path = 'report/report.docx'
    doc.save(file_path)
    return file_path

def generate_conference_list(student_data, tech_data):
    doc = docx.Document()
    set_document_style(doc)

    first_paragraph = doc.add_paragraph(
        'Список представляемых к публикации докладов',
        style='Normal'
    )
    first_paragraph.runs[0].bold = True
    first_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Секция кафедры
    doc.add_paragraph(f"Кафедра {tech_data[0][1]}")
    doc.add_paragraph(tech_data[0][2])
    doc.add_paragraph(f"e-mail: {tech_data[0][4]}")
    doc.add_paragraph(f"тел.: {tech_data[0][5]}")

    # Пополнение списка студентов
    for row in student_data:
        if len(row) > 16 and row[16] == "1" or row[16] == "2":
            combined_paragraph = doc.add_paragraph()
            
            name_run = combined_paragraph.add_run(convert_to_initials(row[7] + " " + row[8] + " " + row[9]))
            name_run.italic = True
            
            combined_paragraph.add_run(f"{row[13]}")

    doc.add_paragraph("\n" * 2)
    doc.add_paragraph(f"Руководитель УНИДС {' ' * 40}{convert_to_initials(tech_data[0][2])}")

    file_path = 'report/publications.docx'
    doc.save(file_path)
    return file_path

@app.get("/conferences/programme")
def get_programme() -> FileResponse:
    tech_data = load_google_sheet(GOOGLE_SHEET_ID, TECH_RANGE)
    student_data = load_google_sheet(GOOGLE_SHEET_ID, STUD_RANGE)
    if not tech_data or not student_data:
        raise HTTPException(status_code=404, detail="Conference data not found")

    # Generate the program document
    file_path = generate_conference_program(student_data, tech_data)

    # Return the FileResponse with the generated file
    return FileResponse(path=file_path, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename="conference_programme.docx")


# Endpoint for generating conference report document
@app.get("/conferences/report")
def get_report() -> FileResponse:
    tech_data = load_google_sheet(GOOGLE_SHEET_ID, TECH_RANGE)
    student_data = load_google_sheet(GOOGLE_SHEET_ID, STUD_RANGE)
    if (not tech_data) or (not student_data):
        raise HTTPException(status_code=404, detail="Conference data not found")

    file_path = generate_conference_report(student_data, tech_data)
    return FileResponse(file_path, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename="conference_report.docx")

# Endpoint for generating conference publications list document
@app.get("/conferences/publications")
def get_publications() -> FileResponse:
    tech_data = load_google_sheet(GOOGLE_SHEET_ID, TECH_RANGE)
    student_data = load_google_sheet(GOOGLE_SHEET_ID, STUD_RANGE)
    if (not tech_data) or (not student_data):
        raise HTTPException(status_code=404, detail="Conference data not found")

    file_path = generate_conference_list(student_data, tech_data)
    return FileResponse(file_path, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename="conference_publications.docx")

if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="127.0.0.1", port=8000)

#uvicorn v4:app --reload --log-level info
#http://127.0.0.1:8000