from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from googleapiclient.discovery import build
from google.oauth2.service_account import Credentials
from datetime import datetime
import docx
import os

# Загрузка данных из Google Sheets
def load_google_sheet(s_id, s_range):
    creds = Credentials.from_service_account_file('service.json')
    service = build('sheets', 'v4', credentials=creds)
    sheet = service.spreadsheets()
    result = sheet.values().get(spreadsheetId=s_id, range=s_range).execute()
    return result.get('values', [])

# Установка стиля для документа
def set_document_style(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.5
    paragraph_format.first_line_indent = Cm(1)
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.sections[0].left_margin = Cm(2)
    doc.sections[0].right_margin = Cm(2)
    doc.sections[0].top_margin = Cm(2)
    doc.sections[0].bottom_margin = Cm(2)

def convert_to_initials(full_name):
    parts = full_name.split()
    if len(parts) == 3:
        surname, name, patronymic = parts
        return f"{surname} {name[0]}.{patronymic[0]}."
    elif len(parts) == 2:
        surname, name = parts
        return f"{surname} {name[0]}."
    else:
        return full_name

def format_date(date_str):
    months = {
        '01': 'января', '02': 'февраля', '03': 'марта', '04': 'апреля',
        '05': 'мая', '06': 'июня', '07': 'июля', '08': 'августа',
        '09': 'сентября', '10': 'октября', '11': 'ноября', '12': 'декабря'
    }
    date = datetime.strptime(date_str, "%Y-%m-%d")
    return f"{date.day} {months[date.strftime('%m')]} {date.year}г."

def generate_conference_program(tech_data):
    doc = docx.Document()
    set_document_style(doc)

    first_paragraph = doc.add_paragraph(
        'Форма представления материалов для программы 78 МСНК ГУАП',
        style='Normal'
    )
    first_paragraph.runs[0].bold = True
    first_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Секция кафедры
    section_heading = doc.add_paragraph()
    run1 = section_heading.add_run(f" {' ' * 4} Секция каф. 43. Компьютерных технологий и программной инженерии")
    run1.bold = True
    run1.italic = True

    # Преобразуем строки с датами в объекты datetime и сортируем
    tech_data.sort(key=lambda x: datetime.strptime(x[15], "%Y-%m-%d"))
    
    # Словарь для хранения студентов по заседаниям
    sessions = {}
    session_num = 1

    for row in tech_data:
        date = row[15]
        if date not in sessions:
            sessions[date] = session_num
            session_num += 1

    for date, session_id in sessions.items():
        session_heading = doc.add_paragraph(f'Заседание {str(session_id)}', style='Normal')
        
        session_info = doc.add_paragraph()
        session_info.add_run(f"{format_date(date)}")

        participant_num = 1
        for row in tech_data:
            if sessions[row[15]] == session_id:
                initials = convert_to_initials(row[7] + " " + row[8] + " " + row[9])  
                doc.add_paragraph(f'{participant_num}. {initials}', style='Normal')
                doc.add_paragraph(f'{row[13]}', style='Normal') 
                participant_num += 1

    doc.save('report/(1) Программа конференции.docx')

def generate_conference_report(tech_data):
    doc = docx.Document()
    set_document_style(doc)

    first_paragraph = doc.add_paragraph(
        'Отчёт о конференции 78 МСНК ГУАП',
        style='Normal'
    )
    first_paragraph.runs[0].bold = True
    first_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    section_heading = doc.add_paragraph()
    run1 = section_heading.add_run(f" {' ' * 4} Секция каф. 43. Компьютерных технологий и программной инженерии ")
    run1.bold = True
    run1.italic = True

    tech_data.sort(key=lambda x: datetime.strptime(x[15], "%Y-%m-%d"))
    
    # Словарь для хранения студентов по заседаниям
    sessions = {}
    session_num = 1

    for row in tech_data:
        date = row[15]
        if date not in sessions:
            sessions[date] = session_num
            session_num += 1

    for date, session_id in sessions.items():
        session_heading = doc.add_paragraph(f'Заседание {str(session_id)}', style='Normal')
        
        session_info = doc.add_paragraph()
        session_info.add_run(f"{format_date(date)}")

        doc.add_paragraph(
            f"Список докладов",
            style='Normal'
        )

        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '№ п/п'
        hdr_cells[1].text = 'ФИО докладчика, название доклада'
        hdr_cells[2].text = 'Статус (магистр/студент)'
        hdr_cells[3].text = 'Решение'

        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.first_line_indent = Cm(0)

        # Заполнение таблицы
        participant_num = 1
        for row in tech_data:
            if sessions[row[15]] == session_id:
                initials = row[7] + " " + row[8] + " " + row[9]
                status = f"{row[12]} Гр. № {row[11]}" if row[11] else row[12]
                
                # Решение с учётом трёх вариантов
                if len(row) > 16:
                    if row[16] == "1":
                        recommendation = "опубликовать доклад в сборнике МСНК"
                    elif row[16] == "2":
                        recommendation = ("опубликовать доклад в сборнике МСНК; "
                                          "рекомендовать к участию в финале конкурса "
                                          "на лучшую студенческую научную работу ГУАП")
                    elif row[16] == "0":
                        recommendation = "доклад плохо подготовлен"
                else:
                    recommendation = "нет данных"

                row_cells = table.add_row().cells
                row_cells[0].text = str(participant_num)
                row_cells[1].text = f"{initials}\n{row[13]}"
                row_cells[2].text = status 
                row_cells[3].text = recommendation

                for paragraph in row_cells[1].paragraphs + row_cells[2].paragraphs + row_cells[3].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    paragraph.paragraph_format.first_line_indent = Cm(0)

                participant_num += 1

        doc.add_paragraph()

    doc.add_paragraph("Подпись научного руководителя секции", style='Normal')

    doc.save('report/(1) Отчёт о конференции.docx')

def generate_conference_list(tech_data):
    doc = docx.Document()
    set_document_style(doc)

    first_paragraph = doc.add_paragraph(
        'Список представляемых к публикации докладов',
        style='Normal'
    )
    first_paragraph.runs[0].bold = True
    first_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Кафедра 43. Компьютерных технологий и программной инженерии")
    
    for row in tech_data:
        if len(row) > 16 and row[16] == "1" or row[16] == "2":
            combined_paragraph = doc.add_paragraph()
            
            name_run = combined_paragraph.add_run(convert_to_initials(row[7] + " " + row[8] + " " + row[9]))
            name_run.italic = True
            
            combined_paragraph.add_run(f"{row[13]}")

    # Добавление строки "Руководитель УНИДС"
    doc.add_paragraph("\n" * 2)
    doc.add_paragraph(f"Руководитель УНИДС {' ' * 40}")

    doc.save('report/(1) Список представляемых к публикации докладов.docx')

if __name__ == "__main__":
    if not os.path.exists('report'):
        os.makedirs('report')

    sheet_id = '1uuvM1XjOtNce025VH5z1PO1Yn02uNpZqu2jII1oRXZo'
    tech_range = 'Sheet1!A2:S'  # Диапазон данных

    tech_data = load_google_sheet(sheet_id, tech_range)
    
    # CLI
    print("Какой документ хотите составить?")
    print("1. Программа конференции")
    print("2. Отчёт о конференции")
    print("3. Список представляемых к публикации докладов")
    print("0. Выйти")
    while True:
        document_type = input("Введите номер документа (1, 2 или 3): ")
        if document_type == '1':
            generate_conference_program(tech_data)
            print("Сгенерирована программа конференции.")
        elif document_type == '2':
            generate_conference_report(tech_data)
            print("Сгенерирован отчет о конференции.")
        elif document_type == '3':
            generate_conference_list(tech_data)
            print("Сгенерирован список представляемых к публикации докладов")
        elif document_type == '0':
            print("Завершение программы")
            break
        else:
            print("Ошибка: Неправильный формат ввода.")
